"""
Ejecutar una sola vez para generar el manual Word:
    pip install python-docx
    python generar_manual.py
"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess, sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

from datetime import date

# ── Paleta de colores ─────────────────────────────────────────────────────────
AZUL_OSCURO  = RGBColor(0x1B, 0x4F, 0x8A)
AZUL_MEDIO   = RGBColor(0x2E, 0x6D, 0xB8)
GRIS_CLARO   = RGBColor(0xF5, 0xF7, 0xFA)
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)
GRIS_TEXTO   = RGBColor(0x44, 0x44, 0x44)

doc = Document()

# ── Márgenes ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Estilos base ─────────────────────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = GRIS_TEXTO

def set_heading_style(para, text, level, color=AZUL_OSCURO, size=14):
    para.clear()
    run = para.add_run(text)
    run.font.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after  = Pt(6)

def add_heading(text, level=1, color=AZUL_OSCURO, size=14):
    p = doc.add_paragraph()
    set_heading_style(p, text, level, color, size)
    return p

def add_body(text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = GRIS_TEXTO
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        rb = p.add_run(bold_prefix + " ")
        rb.font.bold = True
        rb.font.name = "Calibri"
        rb.font.size = Pt(11)
        rb.font.color.rgb = AZUL_MEDIO
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = GRIS_TEXTO
    p.paragraph_format.space_after = Pt(3)
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_module_table(rows_data, col_widths=None):
    table = doc.add_table(rows=1 + len(rows_data), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(["Acción", "Descripción"]):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = BLANCO
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        shade_cell(hdr[i], "1B4F8A")
    for ri, (accion, desc) in enumerate(rows_data, 1):
        cells = table.rows[ri].cells
        cells[0].text = accion
        cells[1].text = desc
        for ci, cell in enumerate(cells):
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell.text)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = GRIS_TEXTO
            if ri % 2 == 0:
                shade_cell(cell, "EBF1F8")
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return table


# ════════════════════════════════════════════════════════════════════════════
# PORTADA
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MANUAL DE USUARIO")
run.font.bold = True
run.font.size = Pt(26)
run.font.color.rgb = AZUL_OSCURO
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Sistema de Gestión Disciplinaria — OCDI")
run.font.bold = True
run.font.size = Pt(16)
run.font.color.rgb = AZUL_MEDIO
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\nSecretaría Distrital de Salud de Bogotá\nOficina de Control Disciplinario Interno")
run.font.size = Pt(12)
run.font.color.rgb = GRIS_TEXTO
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"\n\nVersión 4.2  ·  {date.today().strftime('%B de %Y').capitalize()}")
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = GRIS_TEXTO
run.font.name = "Calibri"

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ════════════════════════════════════════════════════════════════════════════
add_heading("1. Introducción", size=15)
add_body(
    "El Sistema OCDI es la plataforma web oficial de la Oficina de Control Disciplinario Interno de la "
    "Secretaría Distrital de Salud de Bogotá (SDS). Centraliza y digitaliza la gestión de expedientes "
    "disciplinarios, correspondencia, autos, quejas y disponibilidad de sala, reemplazando el manejo "
    "manual en hojas de cálculo Excel."
)
add_body(
    "La aplicación funciona dentro de la red local de la oficina (LAN) y no requiere conexión a internet."
)

# ════════════════════════════════════════════════════════════════════════════
# 2. ACCESO AL SISTEMA
# ════════════════════════════════════════════════════════════════════════════
add_heading("2. Acceso al sistema", size=15)
add_body("Para ingresar a la plataforma:")
add_bullet("Abra cualquier navegador web (Chrome, Edge o Firefox).", bold_prefix="1.")
add_bullet("En la barra de direcciones escriba:  http://[IP_DEL_SERVIDOR]:8000", bold_prefix="2.")
add_bullet("Ingrese su nombre de usuario y contraseña en la pantalla de inicio de sesión.", bold_prefix="3.")
add_bullet("Haga clic en Ingresar.", bold_prefix="4.")
add_body(
    "Si no conoce la dirección IP del servidor o sus credenciales, comuníquese con el administrador del sistema.",
    italic=True
)

add_heading("2.1 Tipos de usuario", size=12, color=AZUL_MEDIO)
add_module_table([
    ("Administrador", "Acceso total: puede crear, editar, eliminar registros y gestionar usuarios del sistema."),
    ("Abogado / Usuario estándar", "Puede registrar y consultar información según los módulos habilitados para su perfil."),
    ("Solo lectura", "Puede consultar información pero no realizar cambios."),
], col_widths=[5, 11])

# ════════════════════════════════════════════════════════════════════════════
# 3. PORTAL PRINCIPAL
# ════════════════════════════════════════════════════════════════════════════
add_heading("3. Portal principal", size=15)
add_body(
    "Al iniciar sesión se muestra el Portal con tarjetas de acceso rápido a cada módulo. "
    "Cada tarjeta indica el número total de registros activos y permite navegar directamente. "
    "También está disponible la opción de descargar el Backup Completo en formato ZIP."
)

# ════════════════════════════════════════════════════════════════════════════
# 4. MÓDULOS DEL SISTEMA
# ════════════════════════════════════════════════════════════════════════════
add_heading("4. Módulos del sistema", size=15)
add_body(
    "La barra lateral izquierda permite acceder a los seis módulos principales. "
    "A continuación se describe cada uno."
)

# -- 4.1 Base Expedientes
add_heading("4.1 Base Expedientes", size=13, color=AZUL_MEDIO)
add_body(
    "Es el módulo central del sistema. Contiene el registro completo de todos los expedientes "
    "disciplinarios con sus datos de identificación, partes involucradas, asunto, etapas procesales "
    "y estado."
)
add_module_table([
    ("Dashboard", "Vista resumen con estadísticas, gráficos por etapa y expedientes próximos a vencer."),
    ("Lista de Expedientes", "Tabla paginada con filtros por año, abogado, etapa, estado y búsqueda libre. Muestra semáforos de vencimiento en color (verde / amarillo / rojo)."),
    ("Nuevo Expediente", "Formulario de 11 secciones para registrar un expediente completo. Incluye selectores interactivos de Tipología y Área de Origen."),
    ("Editar / Ver detalle", "Permite consultar o modificar cualquier expediente existente."),
    ("Importar Excel", "Carga masiva de expedientes desde un archivo Excel con el formato estándar."),
    ("Exportar filtrado", "Descarga un Excel personalizado eligiendo campos y filtros activos."),
    ("Seguimiento Mensual", "Tabla de actuaciones mes a mes por expediente. Haga clic en cualquier celda para registrar o editar la actuación de ese mes."),
], col_widths=[5, 11])

# -- 4.2 Lista de Reparto de Abogados
add_heading("4.2 Lista de Reparto de Abogados (Correspondencia)", size=13, color=AZUL_MEDIO)
add_body(
    "Registra y controla toda la correspondencia recibida: derechos de petición, tutelas, "
    "proposiciones del Consejo, requerimientos de entes de control y comunicaciones internas/externas."
)
add_module_table([
    ("Listar oficios", "Tabla con semáforo dual: verde si ya tiene respuesta, amarillo si está próximo a vencer, rojo si está vencido."),
    ("Nuevo oficio", "Formulario para registrar un nuevo radicado con abogado responsable, término en días, tipo de requerimiento y URL de respuesta."),
    ("Editar / Ver", "Consultar o modificar un oficio existente, incluyendo radicados de salida."),
    ("Importar / Exportar", "Importar desde Excel histórico o exportar la lista con los 19 campos en el formato oficial."),
], col_widths=[5, 11])
add_body(
    "El semáforo considera festivos colombianos y descuenta 2 días hábiles del plazo para dar margen de revisión previa.",
    italic=True
)

# -- 4.3 Control de Autos
add_heading("4.3 Control de Autos de Sustanciación y/o Trámites", size=13, color=AZUL_MEDIO)
add_body(
    "Registro formal de todos los autos producidos en los expedientes. Corresponde al formato "
    "oficial SDS-CDO-FT-001 v4."
)
add_module_table([
    ("Listar autos", "Tabla con filtros por expediente, abogado, asunto y fecha. Bordes y filas alternas para facilitar la lectura."),
    ("Nuevo auto", "Formulario con selección de expediente, número, fecha, asunto predefinido y abogado responsable."),
    ("Exportar", "Descarga en formato Excel conservando la estructura del formato oficial."),
    ("Importar", "Carga masiva desde Excel histórico."),
], col_widths=[5, 11])

# -- 4.4 SDQS
add_heading("4.4 SDQS — Quejas y Solicitudes", size=13, color=AZUL_MEDIO)
add_body(
    "Módulo para gestionar las quejas y solicitudes recibidas a través del Sistema Distrital de "
    "Quejas y Soluciones (SDQS)."
)
add_module_table([
    ("Listar SDQS", "Tabla con filtros por mes, competencia OCDI (SÍ/NO), responsable y búsqueda libre. Indicador visual de COMPETENCIA OCDI."),
    ("Nueva queja/solicitud", "Formulario con 5 secciones. Si se marca COMPETENCIA OCDI = SÍ, aparecen automáticamente los campos BPM y Responsable."),
    ("Editar / Ver detalle", "Consultar o modificar un registro SDQS."),
    ("Exportar", "Descarga en Excel con los 16 campos del registro."),
    ("Importar", "Carga masiva desde el archivo BASE SDQS.xlsx con normalización automática de nombres de responsables."),
], col_widths=[5, 11])

# -- 4.5 Expedientes Digitales
add_heading("4.5 Seguimiento Expedientes Digitales 2025–2026", size=13, color=AZUL_MEDIO)
add_body(
    "Módulo específico para los expedientes digitales del período 2025–2026, con trazabilidad de "
    "comunicaciones, fechas de revisión y alertas de vencimiento."
)
add_module_table([
    ("Lista de expedientes digitales", "Tabla con alertas visuales por proximidad de vencimiento y estado de comunicaciones."),
    ("Nuevo expediente digital", "Formulario con datos del expediente, autos y comunicaciones asociadas."),
    ("Comunicaciones", "Registro de radicados de salida, fechas de envío, seguimiento y respuesta por expediente."),
    ("Exportar", "Descarga en Excel con todas las comunicaciones asociadas a cada expediente."),
], col_widths=[5, 11])

# -- 4.6 Sala de Audiencias
add_heading("4.6 Disponibilidad Horaria — Sala de Audiencias", size=13, color=AZUL_MEDIO)
add_body(
    "Calendario mensual para gestionar la disponibilidad de la sala de audiencias, "
    "organizado en franjas horarias."
)
add_module_table([
    ("Vista mensual", "Calendario con franjas del día. Las celdas muestran si la franja está Disponible u Ocupada con el título del evento."),
    ("Registrar evento", "Haga clic en una franja disponible para agregar título, descripción, responsable y estado del evento."),
    ("Editar / Cancelar", "Modifique o elimine un evento existente."),
    ("Exportar", "Descarga en Excel con todos los eventos del mes o período seleccionado."),
], col_widths=[5, 11])

# ════════════════════════════════════════════════════════════════════════════
# 5. BACKUP Y RESPALDO
# ════════════════════════════════════════════════════════════════════════════
add_heading("5. Respaldo de información (Backup)", size=15)
add_body(
    "Desde el Portal principal o desde el menú Backup puede descargar dos tipos de respaldo:"
)
add_module_table([
    ("Excel Completo (.xlsx)", "Un solo archivo Excel con 4 hojas: Base Expedientes, Expedientes Digitales, Sala de Audiencias y Control de Autos."),
    ("Backup Completo (.zip)", "Archivo ZIP con 6 carpetas, una por cada módulo. La carpeta Base Expedientes incluye una segunda hoja con el Seguimiento Mensual."),
], col_widths=[5, 11])
add_body(
    "Se recomienda realizar el backup completo al menos una vez a la semana y guardarlo en un "
    "lugar seguro fuera del servidor.",
    italic=True
)

# ════════════════════════════════════════════════════════════════════════════
# 6. SEMÁFOROS Y ALERTAS
# ════════════════════════════════════════════════════════════════════════════
add_heading("6. Semáforos y alertas", size=15)
add_body(
    "El sistema utiliza indicadores de color en varias vistas para facilitar el seguimiento de plazos:"
)
add_module_table([
    ("🟢 Verde", "Dentro del plazo / respondido."),
    ("🟡 Amarillo", "Próximo a vencer (entre 6 y 8 días transcurridos sin respuesta)."),
    ("🔴 Rojo", "Vencido (9 o más días sin respuesta, o fecha de vencimiento superada)."),
    ("⚪ Gris / sin color", "Sin término definido o registro nuevo."),
], col_widths=[5, 11])

# ════════════════════════════════════════════════════════════════════════════
# 7. RECOMENDACIONES
# ════════════════════════════════════════════════════════════════════════════
add_heading("7. Recomendaciones de uso", size=15)
add_bullet("No comparta sus credenciales de acceso con otras personas.")
add_bullet("Cierre sesión al finalizar su jornada (botón Salir en la barra lateral).")
add_bullet("Ante cualquier error o comportamiento inesperado, tome una captura de pantalla y repórtelo al administrador.")
add_bullet("Realice el backup ZIP periódicamente como copia de seguridad.")
add_bullet("Los campos obligatorios están marcados con asterisco (*) en los formularios.")
add_bullet("Los cambios en Seguimiento Mensual se guardan celda por celda — no hay botón de guardar global.")

# ════════════════════════════════════════════════════════════════════════════
# 8. CONTACTO
# ════════════════════════════════════════════════════════════════════════════
add_heading("8. Soporte y contacto", size=15)
add_body(
    "Para reportar fallas técnicas, solicitar nuevos accesos o capacitación, comuníquese con:"
)
add_bullet("Administrador del sistema: [Nombre del administrador]", bold_prefix="▸")
add_bullet("Correo: [correo@saludcapital.gov.co]", bold_prefix="▸")
add_bullet("Extensión: [número de extensión]", bold_prefix="▸")

# ── Pie de página ─────────────────────────────────────────────────────────────
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run(
    f"Sistema OCDI — Secretaría Distrital de Salud de Bogotá  |  Versión 4.2  |  {date.today().strftime('%Y')}"
)
run.font.size = Pt(9)
run.font.color.rgb = GRIS_TEXTO
run.font.name = "Calibri"
run.font.italic = True

# ── Guardar ───────────────────────────────────────────────────────────────────
nombre = f"Manual_Usuario_OCDI_v4.2_{date.today().strftime('%Y%m%d')}.docx"
doc.save(nombre)
print(f"Manual generado: {nombre}")
